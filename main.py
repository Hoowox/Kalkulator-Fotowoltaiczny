from tkinter import ttk
from tkinter import messagebox
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure
import pandas as pd
import glob
from tkinter import filedialog
import json
import numpy as np
from mpl_toolkits.mplot3d.art3d import Poly3DCollection
from mpl_toolkits.mplot3d import Axes3D
from PIL import Image, ImageTk
import io
import os
import sys
import tkinter as tk
from docx import Document

"""
Uwagi do programu:
- matplotlib mial problemy z przenikaniem się warstw wizualizacji 3D stad kod jest troche zlozoy
- w programie moga znajdowac sie zmienne dotyczace naswietlen, ktore w rzeczywistosci odpowiadaja za liczbe godzin slonecznych
- program przy kalkulacji nie korzysta z wszystkich danych ze wzgledu na trudnosc w ich odpowiednim przeliczeniu
- do dzialania programu niezbedne poza bibliotekami sa pliki polski.json oraz english.json
"""

class CustomDialog(tk.Toplevel):
    def __init__(self, parent, title=None, message=None, initialvalue=None, texts=None):
        tk.Toplevel.__init__(self, parent)
        self.title(title)
        self.geometry('200x120')
        self.resizable(False, False)
        self.protocol("WM_DELETE_WINDOW", self.cancel)
        self.result = None
        self.texts = texts

        # Zastosowanie stylu 'TLabel' dla etykiety
        label = ttk.Label(self, text=message)
        label.pack(padx=10, pady=10)

        # Zastosowanie stylu 'TEntry' dla pola wprowadzania
        self.entry = ttk.Entry(self)
        self.entry.pack(fill='both', expand=True, padx=20, pady=5)  # Zmniejszenie pola do wpisywania
        self.entry.insert(0, initialvalue)

        buttons = ttk.Frame(self)
        buttons.pack(pady=5)  # Dodanie odstępu od dolnej krawędzi okna

        # Zastosowanie stylu 'TButton' dla przycisków
        ok_button = ttk.Button(buttons, text=self.texts['save'], command=self.ok, width=10)
        ok_button.pack(side='left', padx=5, pady=5)  # Zwiększenie wysokości przycisku

        cancel_button = ttk.Button(buttons, text=self.texts['cancel_settings'], command=self.cancel, width=10)
        cancel_button.pack(side='right', padx=5, pady=5)  # Zwiększenie wysokości przycisku

    def ok(self, event=None):
        self.result = self.entry.get()
        self.destroy()

    def cancel(self, event=None):
        self.destroy()

def askstring(title, message, initialvalue=None, texts=None):
    root = tk.Tk()
    root.withdraw()
    d = CustomDialog(root, title, message, initialvalue, texts)
    root.wait_window(d)
    return d.result

class MojaAplikacja:
    def __init__(self, root):
        self.root = root
        self.root.state('zoomed')  # Działa w Windows, dla Linux użyj: self.root.attributes('-zoomed', True)
        self.root.bind('<F11>', self.toggle_fullscreen)
        self.root.bind('<Escape>', lambda e: self.root.destroy())
        self.naswietlenie_miesieczne = {}
        self.temperatura_miesieczna = {}
        self.language = 'polski'
        self.texts = self.load_language(self.language)
        self.current_view = None
        self.auto_rotate_id = None
        self.event_handlers = []
        self.help_images = []  # Przechowuje referencje do obrazów

        # Ustaw ikonę okna
        try:
            # Pobierz absolutną ścieżkę do folderu z main.py
            base_dir = os.path.dirname(os.path.abspath(__file__))

            # Zbuduj pełną ścieżkę do ikony
            icon_path = self.get_resource_path("ikona.ico")

            # Ustaw ikonę
            root.iconbitmap(icon_path)
        except Exception as e:
            print(f"Nie udało się załadować ikony: {e}")
            print(f"Szukana ścieżka: {icon_path}")  # Debugowanie ścieżki

        self.root.title(self.texts['title'])
        self.root.bind('<F1>', self.pokaz_pomoc)

        # Tworzenie paska menu
        self.menu_bar = tk.Menu(root)
        self.root.config(menu=self.menu_bar)

        # Menu ustawień
        self.settings_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label=self.texts['settings'], menu=self.settings_menu)
        self.settings_menu.add_command(label=self.texts['monthly sunshine hours'], command=self.ustaw_naswietlenie)
        self.settings_menu.add_command(label=self.texts['monthly_temperatures'], command=self.ustaw_temperatury)
        language_menu = tk.Menu(self.settings_menu, tearoff=0)
        self.settings_menu.add_cascade(label=self.texts['change_language'], menu=language_menu)
        language_menu.add_command(label="Polski", command=lambda: self.change_language('polski'))
        language_menu.add_command(label="English", command=lambda: self.change_language('english'))

        # Menu widoków
        self.view_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label=self.texts['view'], menu=self.view_menu)
        self.view_menu.add_command(label=self.texts['bar_chart'], command=self.pokaz_wykres_slupkowy)
        self.view_menu.add_command(label=self.texts['3D_visualization'], command=self.pokaz_wizualizacje_3d)

        # Menu zapisywania
        self.save_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label=self.texts['save'], menu=self.save_menu)
        self.save_menu.add_command(label=self.texts['save_chart'], command=self.zapisz_wykres)
        self.save_menu.add_command(label=self.texts['save_settings'], command=self.zapisz_ustawienia)
        self.save_menu.add_command(label=self.texts['save_data'], command=self.zapisz_dane)

        # Menu importu
        self.import_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label=self.texts['import'], menu=self.import_menu)
        self.import_menu.add_command(label=self.texts['import_chart'], command=self.importuj_wykres)
        self.import_menu.add_command(label=self.texts['import_settings'], command=self.wybierz_i_importuj_ustawienia)
        self.import_menu.add_command(label=self.texts['import_data'], command=self.wybierz_i_importuj_dane)

        # Menu pomocy
        self.help_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label=self.texts['help'], menu=self.help_menu)
        self.help_menu.add_command(label=self.texts['how_to_use'], command=self.pokaz_pomoc)

        # Dodawanie opcji Inne
        other_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label=self.texts['other'], menu=other_menu)

        # Dodawanie opcji wyjścia z programu
        self.exit_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label=self.texts['exit'], menu=self.exit_menu)
        self.exit_menu.add_command(label=self.texts['exit'], command=self.wyjdz)

        # Kontener na wykresy i pomoc
        self.chart_container = tk.Frame(root)
        self.chart_container.grid(row=0, column=1, sticky='nsew')


        # Ramka dla wykresów
        self.chart_frame = tk.Frame(self.chart_container)
        self.chart_frame.pack(fill='both', expand=True)

        # Ramka dla pomocy (początkowo ukryta)
        self.help_frame = tk.Frame(self.chart_container)

        # Inicjalizacja wykresu
        self.fig = Figure(figsize=(6, 6), dpi=100)
        self.canvas = FigureCanvasTkAgg(self.fig, master=self.chart_frame)
        self.canvas.get_tk_widget().pack(fill='both', expand=True)

        # Inicjalizacja pustych osi (do dodania później w zależności od widoku)
        self.ax = None

        # Rozciągnięcie kolumny z wykresem
        root.grid_columnconfigure(1, weight=1)
        root.grid_rowconfigure(0, weight=1)

        # Tworzenie ramki na ustawienia wejściowe
        input_frame = ttk.Frame(root)
        input_frame.grid(row=0, column=0, padx=10, pady=10, sticky='ns')

        # Etykieta wyników
        self.result_label = tk.Label(root, text="")
        self.result_label.grid(row=1, column=0, padx=10, pady=10, sticky='ns')
        # Etykiety i pola do wprowadzania danych
        labels = ["Liczba paneli", "Kąt nachylenia (°)",
                  "Azymut (°)", "Moc jednego panelu (W)", "Sprawność paneli (%)",
                  "Temperatura otoczenia (°C)", "Liczba godzin słonecznych", "Albedo (%)",
                  "Zacienienie (%)", "Zanieczyszczenia (%)"]

        self.labels = {}
        self.entries = {}
        for i, label in enumerate(labels):
            lbl = ttk.Label(input_frame, text=label)
            lbl.grid(row=i, column=0, sticky='w')
            self.labels[label] = lbl
            entry = ttk.Entry(input_frame)
            entry.grid(row=i, column=1)
            self.entries[label] = entry

        # Przycisk do obliczeń
        self.calculate_button = ttk.Button(input_frame, text=self.texts['calculate'], command=self.oblicz)
        self.calculate_button.grid(row=len(labels), column=0, columnspan=2, pady=10)

        # Etykieta na wynik
        self.result_label = ttk.Label(root, text=self.texts['result'])
        self.result_label.grid(row=1, column=0, padx=10, pady=10)

    def get_resource_path(self, relative_path):
        """Pobiera ścieżkę do zasobów w wersji .exe i developerskiej"""
        if getattr(sys, 'frozen', False):
            base_path = sys._MEIPASS
        else:
            base_path = os.path.dirname(os.path.abspath(__file__))

        return os.path.join(base_path, relative_path)

    def toggle_fullscreen(self, event=None):
        is_fullscreen = self.root.attributes('-fullscreen')
        self.root.attributes('-fullscreen', not is_fullscreen)

    def load_language(self, lang):
        # Określ ścieżkę do pliku .json w zależności od trybu działania (dev vs .exe)
        if getattr(sys, 'frozen', False):
            base_path = sys._MEIPASS  # Tryb .exe
        else:
            base_path = os.path.dirname(os.path.abspath(__file__))  # Tryb developerski

        file_path = os.path.join(base_path, f"{lang}.json")

        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    def update_ui_texts(self):
        # Aktualizacja tekstów w menu
        self.settings_menu.entryconfig(0, label=self.texts['monthly sunshine hours'])
        self.settings_menu.entryconfig(1, label=self.texts['monthly_temperatures'])
        self.settings_menu.entryconfig(2, label=self.texts['change_language'])
        self.view_menu.entryconfig(0, label=self.texts['bar_chart'])
        self.view_menu.entryconfig(1, label=self.texts['3D_visualization'])
        self.save_menu.entryconfig(0, label=self.texts['save_chart'])
        self.save_menu.entryconfig(1, label=self.texts['save_settings'])
        self.save_menu.entryconfig(2, label=self.texts['save_data'])
        self.import_menu.entryconfig(0, label=self.texts['import_chart'])
        self.import_menu.entryconfig(1, label=self.texts['import_settings'])
        self.import_menu.entryconfig(2, label=self.texts['import_data'])
        self.help_menu.entryconfig(0, label=self.texts['how_to_use'])
        self.exit_menu.entryconfig(0, label=self.texts['exit'])
        self.menu_bar.entryconfig(1, label=self.texts['settings'])
        self.menu_bar.entryconfig(2, label=self.texts['view'])
        self.menu_bar.entryconfig(3, label=self.texts['save'])
        self.menu_bar.entryconfig(4, label=self.texts['import'])
        self.menu_bar.entryconfig(5, label=self.texts['help'])
        self.menu_bar.entryconfig(6, label=self.texts['other'])
        self.menu_bar.entryconfig(7, label=self.texts['exit'])

        # Aktualizacja tekstu przycisku i etykiety wyniku
        self.calculate_button['text'] = self.texts['calculate']
        self.result_label['text'] = self.texts['result']

        # Aktualizacja tytułu okna
        self.root.title(self.texts.get('title', ''))

        # Aktualizacja etykiet pól wprowadzania danych
        for label, lbl in self.labels.items():
            lbl['text'] = self.texts.get(label, label)

        # Aktualizacja tytułów okien temperatury i naświetlenia, jeśli są otwarte
        if hasattr(self, 'temperatura_window') and self.temperatura_window.winfo_exists():
            self.temperatura_window.title(self.texts['monthly_temperatures'])
        if hasattr(self, 'naswietlenie_window') and self.naswietlenie_window.winfo_exists():
            self.naswietlenie_window.title(self.texts['monthly sunshine hours'])

        # Aktualizacja komunikatu o błędzie przy wprowadzaniu złych danych
        if hasattr(self, 'result_label') and self.result_label['text'] == "Proszę wprowadzić temperaturę otoczenia.":
            self.result_label['text'] = self.texts['enter_environment_temperature']

        # Aktualizacja etykiety osi Y na wykresie
        if hasattr(self, 'ax'):
            self.ax.set_ylabel(self.texts['generated_energy_kWh'])

        # Aktualizacja komunikatu o usuwaniu naświetleń oraz temperatur
        if hasattr(self, 'naswietlenie_window') and self.naswietlenie_window.winfo_exists():
            self.naswietlenie_window.title(self.texts['delete_all_data'])

        if hasattr(self, 'temperatura_window') and self.temperatura_window.winfo_exists():
            self.temperatura_window.title(self.texts['delete_all_data'])

        # Aktualizacja tekstów związanych z zapisywaniem i importem
        self.file_name_prompt = self.texts['file_name_prompt']
        self.enter_file_name = self.texts['enter_file_name']
        self.info = self.texts['info']
        self.chart_saved = self.texts['chart_saved']
        self.settings_saved = self.texts['settings_saved']
        self.data_saved = self.texts['data_saved']
        self.chart_imported = self.texts['chart_imported']
        self.wrong_number_of_columns = self.texts['wrong_number_of_columns']
        self.missing_expected_column = self.texts['missing_expected_column']
        self.select_and_import_settings = self.texts['select_and_import_settings']
        self.select_and_import_data = self.texts['select_and_import_data']
        self.in_this_folder = self.texts['in_this_folder']

        # Aktualizacja tytułu okna pomocy
        self.help_title = self.texts['help']

        # Aktualizacja tekstu wyświetlanego, gdy plik pomocy nie zostanie znaleziony
        self.help_not_found = self.texts['help_not_found']

        # Aktualizacja tekstu wyświetlanego w oknie dialogowym przy wyjściu
        self.exit_text = self.texts['exit']
        self.exit_confirmation = self.texts['exit_confirmation']

        # Aktualizacja tekstu związanych z oknem dialogowym
        self.save = self.texts['save']
        self.cancel_settings = self.texts['cancel_settings']

    def change_language(self, lang):
        self.language = lang
        self.texts = self.load_language(self.language)
        self.update_ui_texts()

    def set_view_to_slupkowy(self):
        self.current_view = 'Wykres słupkowy'
        self.update_view()

    def set_view_to_3d(self):
        self.current_view = 'Wizualizacja 3D'
        self.update_view()

    def pokaz_wykres_slupkowy(self):
        # Sprawdź, czy help_frame istnieje, i ukryj go
        if hasattr(self, 'help_frame'):
            self.help_frame.pack_forget()

        # Odbinduj zdarzenia scrollowania
        self.root.unbind_all("<MouseWheel>")
        self.root.unbind_all("<Button-4>")
        self.root.unbind_all("<Button-5>")
        self.help_frame.pack_forget()
        self.chart_frame.pack(fill='both', expand=True)
        self.clear_current_view()
        try:
            # Upewnij się, że dane są dostępne
            if not hasattr(self, 'miesiace') or not hasattr(self, 'energia_miesieczna'):
                print("Dane do wykresu nie są dostępne.")
                return

            # Tworzenie osi wykresu słupkowego
            if self.ax is not None:
                self.ax.clear()
            else:
                self.ax = self.fig.add_subplot(111)

            bars = self.ax.bar(range(1, 13), self.energia_miesieczna)
            self.ax.set_ylabel(self.texts['generated_energy_kWh'])
            self.ax.set_xticks(range(1, 13))
            self.ax.set_xticklabels(self.miesiace)

            # Dodawanie liczbowych etykiet nad każdym słupkiem
            for bar, energia in zip(bars, self.energia_miesieczna):
                height = bar.get_height()
                self.ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    height,
                    f'{energia:.2f}',
                    ha='center', va='bottom'
                )

            self.canvas.draw()
        except Exception as e:
            print(f"Błąd podczas tworzenia wykresu: {e}")

    def pokaz_wizualizacje_3d(self):
        # Sprawdź, czy help_frame istnieje, i ukryj go
        if hasattr(self, 'help_frame'):
            self.help_frame.pack_forget()

        # Odbinduj zdarzenia scrollowania
        self.root.unbind_all("<MouseWheel>")
        self.root.unbind_all("<Button-4>")
        self.root.unbind_all("<Button-5>")
        self.help_frame.pack_forget()
        self.chart_frame.pack(fill='both', expand=True)
        self.clear_current_view()

        # Tworzenie figury tylko jeśli nie istnieje
        if self.fig is None:
            self.fig = Figure()

        # Tworzenie płótna tylko jeśli nie istnieje lub zostało usunięte
        if self.canvas is None or self.canvas.get_tk_widget() is None:
            self.canvas = FigureCanvasTkAgg(self.fig, master=self.root)
            self.canvas.get_tk_widget().pack()

        # Tworzenie osi wizualizacji 3D
        self.ax = self.fig.add_subplot(111, projection='3d')
        self.ax.set_axis_off()  # Ukrycie osi

        # Zmienna do kontrolowania kąta automatycznego obrotu
        self.current_azim = 30
        self.current_elev = 30
        self.auto_rotate_speed = 1.5
        self.is_rotating_manually = False
        self.auto_rotate_delay = 2000

        # Sprawdzenie, czy fig.canvas istnieje przed dodaniem eventów
        if self.fig.canvas:
            self.fig.canvas.mpl_connect('motion_notify_event', self.on_mouse_move)
            self.fig.canvas.mpl_connect('button_press_event', self.on_mouse_press)
            self.fig.canvas.mpl_connect('button_release_event', self.on_mouse_release)

        # Automatyczny obrót co 100 ms
        self.auto_rotate_id = self.root.after(100, self.auto_rotate)

        # Wymiary domu
        house_width = 20.0
        house_depth = 15.0
        house_height = 15.0

        # Wymiary dachu
        roof_angle = 35
        roof_angle_rad = np.radians(roof_angle)
        roof_height = (house_width / 2) * np.tan(roof_angle_rad)

        # Definicja okien
        # Wymiary okien
        window_width = 4.0
        window_height = 4.0
        window_thickness = 0.2
        overlap = 0.1

        ### Definicja drzwi ###
        door_width = 4.0
        door_height = window_height + (
                    (house_height - window_height) / 2)
        door_thickness = 0.2

        # Obliczenie pozycji Z dla okien (środek ściany)
        window_z_start = (house_height - window_height) / 2

        # Pozycje okien na przedniej ścianie
        front_window_positions = [
            (house_width / 4 - window_width / 2, window_z_start),
            (3 * house_width / 4 - window_width / 2, window_z_start)
        ]

        # Pozycje okien na bocznych ścianach
        side_window_positions = [
            (house_depth / 2 - window_width / 2, window_z_start)
        ]

        # Pozycje okien i drzwi na tylnej ścianie
        back_openings = [
            (house_width / 4 - window_width / 2, window_z_start, window_height),
            (3 * house_width / 4 - door_width / 2, 0, door_height)
        ]

        ### Tworzenie ścian domu z dziurami na okna i drzwi ###
        # Lista ścian do rysowania
        walls = []

        # Funkcja tworząca dziury w ścianach
        def create_wall_with_holes(wall_start, wall_end, wall_height, openings, axis='x'):
            wall_polygons = []
            positions = sorted(openings, key=lambda x: x[0])
            last_pos = wall_start
            for pos, z_start, height in positions:
                if height == door_height:
                    opening_width = door_width
                else:
                    opening_width = window_width
                pos_end = pos + opening_width

                # Lewa część ściany
                if pos > last_pos:
                    polygon = [
                        [wall_start if axis == 'x' else last_pos, wall_start if axis == 'y' else last_pos, 0],
                        [wall_start if axis == 'x' else pos, wall_start if axis == 'y' else pos, 0],
                        [wall_start if axis == 'x' else pos, wall_start if axis == 'y' else pos, wall_height],
                        [wall_start if axis == 'x' else last_pos, wall_start if axis == 'y' else last_pos, wall_height],
                    ]
                    wall_polygons.append(polygon)
                # Obszar nad otworem
                if z_start + height < wall_height:
                    polygon = [
                        [wall_start if axis == 'x' else pos, wall_start if axis == 'y' else pos, z_start + height],
                        [wall_start if axis == 'x' else pos_end, wall_start if axis == 'y' else pos_end,
                         z_start + height],
                        [wall_start if axis == 'x' else pos_end, wall_start if axis == 'y' else pos_end, wall_height],
                        [wall_start if axis == 'x' else pos, wall_start if axis == 'y' else pos, wall_height],
                    ]
                    wall_polygons.append(polygon)
                # Obszar pod otworem
                if z_start > 0:
                    polygon = [
                        [wall_start if axis == 'x' else pos, wall_start if axis == 'y' else pos, 0],
                        [wall_start if axis == 'x' else pos_end, wall_start if axis == 'y' else pos_end, 0],
                        [wall_start if axis == 'x' else pos_end, wall_start if axis == 'y' else pos_end, z_start],
                        [wall_start if axis == 'x' else pos, wall_start if axis == 'y' else pos, z_start],
                    ]
                    wall_polygons.append(polygon)
                last_pos = pos_end
            # Prawa część ściany
            if last_pos < wall_end:
                polygon = [
                    [wall_start if axis == 'x' else last_pos, wall_start if axis == 'y' else last_pos, 0],
                    [wall_start if axis == 'x' else wall_end, wall_start if axis == 'y' else wall_end, 0],
                    [wall_start if axis == 'x' else wall_end, wall_start if axis == 'y' else wall_end, wall_height],
                    [wall_start if axis == 'x' else last_pos, wall_start if axis == 'y' else last_pos, wall_height],
                ]
                wall_polygons.append(polygon)
            return wall_polygons

        # Przednia ściana z dziurami
        front_wall_polygons = create_wall_with_holes(
            wall_start=0,
            wall_end=house_width,
            wall_height=house_height,
            openings=[(pos, z_start, window_height) for pos, z_start in front_window_positions],
            axis='y'  # Przednia ściana wzdłuż osi Y
        )

        # Lewa ściana z dziurami
        left_wall_polygons = create_wall_with_holes(
            wall_start=0,
            wall_end=house_depth,
            wall_height=house_height,
            openings=[(pos, z_start, window_height) for pos, z_start in side_window_positions],
            axis='x'  # Lewa ściana wzdłuż osi X
        )

        # Prawa ściana z dziurami (kopiujemy pozycje z lewej ściany)
        right_wall_polygons = create_wall_with_holes(
            wall_start=0,
            wall_end=house_depth,
            wall_height=house_height,
            openings=[(pos, z_start, window_height) for pos, z_start in side_window_positions],
            axis='x'  # Prawa ściana wzdłuż osi X
        )

        # Przesunięcie prawej ściany na odpowiednią pozycję
        for polygon in right_wall_polygons:
            for vertex in polygon:
                vertex[0] += house_width

        # Tylna ściana z dziurami na okno i drzwi
        back_wall_polygons = create_wall_with_holes(
            wall_start=0,
            wall_end=house_width,
            wall_height=house_height,
            openings=back_openings,
            axis='y'
        )

        # Przesunięcie tylnej ściany na odpowiednią pozycję
        for polygon in back_wall_polygons:
            for vertex in polygon:
                vertex[1] += house_depth

        # Dodawanie wszystkich ścian do listy
        walls.extend([
            (front_wall_polygons, '#F4A460'),
            (left_wall_polygons, '#BC8F8F'),
            (right_wall_polygons, '#DEB887'),
            (back_wall_polygons, '#D2B48C')
        ])

        # Rysowanie ścian
        for wall_polygons, color in walls:
            for polygon in wall_polygons:
                self.ax.add_collection3d(Poly3DCollection([polygon],
                                                          facecolors=color,
                                                          alpha=1.0))

        # Dodanie okien i drzwi jako niebieskich trójwymiarowych obiektów
        def create_window(window_pos, z_start, wall_pos, axis='x', window_height=4.0, window_width=4.0):
            pos = window_pos - overlap
            pos_end = window_pos + window_width + overlap
            z_end = z_start + window_height + overlap
            wall_thickness = window_thickness

            if axis == 'x':
                vertices = np.array([
                    [wall_pos - wall_thickness, pos, z_start - overlap],
                    [wall_pos + overlap, pos, z_start - overlap],
                    [wall_pos + overlap, pos_end, z_start - overlap],
                    [wall_pos - wall_thickness, pos_end, z_start - overlap],
                    [wall_pos - wall_thickness, pos, z_end],
                    [wall_pos + overlap, pos, z_end],
                    [wall_pos + overlap, pos_end, z_end],
                    [wall_pos - wall_thickness, pos_end, z_end]
                ])
            else:
                vertices = np.array([
                    [pos, wall_pos - wall_thickness, z_start - overlap],
                    [pos_end, wall_pos - wall_thickness, z_start - overlap],
                    [pos_end, wall_pos + overlap, z_start - overlap],
                    [pos, wall_pos + overlap, z_start - overlap],
                    [pos, wall_pos - wall_thickness, z_end],
                    [pos_end, wall_pos - wall_thickness, z_end],
                    [pos_end, wall_pos + overlap, z_end],
                    [pos, wall_pos + overlap, z_end]
                ])
            return vertices

        # Okna na przedniej ścianie
        for x_start, z_start in front_window_positions:
            window_vertices = create_window(
                window_pos=x_start,
                z_start=z_start - overlap,
                wall_pos=0,
                axis='y'
            )
            window_faces = [
                [window_vertices[j] for j in [0, 1, 5, 4]],
                [window_vertices[j] for j in [1, 2, 6, 5]],
                [window_vertices[j] for j in [2, 3, 7, 6]],
                [window_vertices[j] for j in [3, 0, 4, 7]],
                [window_vertices[j] for j in [0, 1, 2, 3]],
                [window_vertices[j] for j in [4, 5, 6, 7]]
            ]
            self.ax.add_collection3d(
                Poly3DCollection(window_faces, facecolors='lightblue', linewidths=1, edgecolors='skyblue', alpha=0.8))

        # Okna na lewej ścianie
        for y_start, z_start in side_window_positions:
            window_vertices = create_window(
                window_pos=y_start,
                z_start=z_start - overlap,
                wall_pos=0,
                axis='x'
            )
            window_faces = [
                [window_vertices[j] for j in [0, 1, 5, 4]],
                [window_vertices[j] for j in [1, 2, 6, 5]],
                [window_vertices[j] for j in [2, 3, 7, 6]],
                [window_vertices[j] for j in [3, 0, 4, 7]],
                [window_vertices[j] for j in [0, 1, 2, 3]],
                [window_vertices[j] for j in [4, 5, 6, 7]]
            ]
            self.ax.add_collection3d(
                Poly3DCollection(window_faces, facecolors='lightblue', linewidths=1, edgecolors='skyblue', alpha=0.8))

        # Okna na prawej ścianie
        for y_start, z_start in side_window_positions:
            window_vertices = create_window(
                window_pos=y_start,
                z_start=z_start - overlap,
                wall_pos=house_width,
                axis='x'
            )
            window_faces = [
                [window_vertices[j] for j in [0, 1, 5, 4]],
                [window_vertices[j] for j in [1, 2, 6, 5]],
                [window_vertices[j] for j in [2, 3, 7, 6]],
                [window_vertices[j] for j in [3, 0, 4, 7]],
                [window_vertices[j] for j in [0, 1, 2, 3]],
                [window_vertices[j] for j in [4, 5, 6, 7]]
            ]
            self.ax.add_collection3d(
                Poly3DCollection(window_faces, facecolors='lightblue', linewidths=1, edgecolors='skyblue', alpha=0.8))

        # Okno i drzwi na tylnej ścianie
        for x_start, z_start, height in [(house_width / 4 - window_width / 2, window_z_start, window_height),
                                         (3 * house_width / 4 - door_width / 2, 0, door_height)]:
            if height == door_height:
                door_vertices = create_window(
                    window_pos=x_start,
                    z_start=z_start,
                    wall_pos=house_depth,
                    axis='y',
                    window_height=door_height,
                    window_width=door_width
                )
                door_faces = [
                    [door_vertices[j] for j in [0, 1, 5, 4]],
                    [door_vertices[j] for j in [1, 2, 6, 5]],
                    [door_vertices[j] for j in [2, 3, 7, 6]],
                    [door_vertices[j] for j in [3, 0, 4, 7]],
                    [door_vertices[j] for j in [0, 1, 2, 3]],
                    [door_vertices[j] for j in [4, 5, 6, 7]]
                ]
                self.ax.add_collection3d(
                    Poly3DCollection(door_faces, facecolors='saddlebrown', linewidths=1, edgecolors='rosybrown', alpha=1.0))

            else:
                # Okno
                window_vertices = create_window(
                    window_pos=x_start,
                    z_start=z_start - overlap,
                    wall_pos=house_depth,
                    axis='y'
                )
                window_faces = [
                    [window_vertices[j] for j in [0, 1, 5, 4]],
                    [window_vertices[j] for j in [1, 2, 6, 5]],
                    [window_vertices[j] for j in [2, 3, 7, 6]],
                    [window_vertices[j] for j in [3, 0, 4, 7]],
                    [window_vertices[j] for j in [0, 1, 2, 3]],
                    [window_vertices[j] for j in [4, 5, 6, 7]]
                ]
                self.ax.add_collection3d(
                    Poly3DCollection(window_faces, facecolors='lightblue', linewidths=1, edgecolors='skyblue', alpha=0.8))

        # Dodanie panelu fotowoltaicznego
        panel_width = house_width / 4
        panel_height = house_depth / 3
        panel_thickness = 1.0
        roof_clearance = 0.9

        # Pozycjonowanie panelu na lewej połaci dachu
        panel_offset_from_ridge = house_width / 4
        panel_x_center = house_width / 2 - panel_offset_from_ridge
        panel_y_center = house_depth / 2

        # Funkcja obliczająca wysokość Z dla dachu w zależności od X
        def roof_z(x):
            return house_height + roof_height - abs(x - house_width / 2) * np.tan(roof_angle_rad)

        # Wierzchołki panelu fotowoltaicznego
        panel_vertices_front = np.array([
            [panel_x_center - panel_width / 2, panel_y_center - panel_height / 2],
            [panel_x_center + panel_width / 2, panel_y_center - panel_height / 2],
            [panel_x_center + panel_width / 2, panel_y_center + panel_height / 2],
            [panel_x_center - panel_width / 2, panel_y_center + panel_height / 2]
        ])

        # Dodanie wysokości Z z uwzględnieniem nachylenia dachu
        panel_vertices_front = np.array([
            [x, y, roof_z(x) + roof_clearance]
            for x, y in panel_vertices_front
        ])

        # Wierzchołki tylnej strony panelu
        panel_vertices_back = np.array([
            [vertex[0], vertex[1], vertex[2] - panel_thickness]
            for vertex in panel_vertices_front
        ])

        # Tworzenie powierzchni panelu
        panel_faces = [
            [panel_vertices_front[j] for j in [0, 1, 2, 3]],
            [panel_vertices_back[j] for j in [0, 1, 2, 3]],
            [panel_vertices_front[0], panel_vertices_front[1], panel_vertices_back[1], panel_vertices_back[0]],
            [panel_vertices_front[1], panel_vertices_front[2], panel_vertices_back[2], panel_vertices_back[1]],
            [panel_vertices_front[2], panel_vertices_front[3], panel_vertices_back[3], panel_vertices_back[2]],
            [panel_vertices_front[3], panel_vertices_front[0], panel_vertices_back[0], panel_vertices_back[3]]
        ]

        # Dodanie panelu jako kolekcji wielokątów
        for face in panel_faces:
            self.ax.add_collection3d(Poly3DCollection([face],
                                                      facecolors='blue', linewidths=1, edgecolors='dodgerblue', alpha=1.0))



        # Rysowanie dachu z dziurą pod panelem
        # Obliczenie pozycji otworu w dachu
        panel_x_start = panel_x_center - panel_width / 2
        panel_x_end = panel_x_center + panel_width / 2
        panel_y_start = panel_y_center - panel_height / 2
        panel_y_end = panel_y_center + panel_height / 2

        # Funkcja obliczająca wysokość Z dla dachu w zależności od X
        def roof_z(x):
            return house_height + roof_height - abs(x - house_width / 2) * np.tan(roof_angle_rad)

        # Lista przechowująca wielokąty lewej połaci dachu (z dziurą)
        left_roof_polygons = []

        # Część lewej połaci powyżej otworu
        left_roof_polygons.append([
            [panel_x_end, panel_y_start, roof_z(panel_x_end)],
            [panel_x_end, panel_y_end, roof_z(panel_x_end)],
            [house_width / 2, panel_y_end, roof_z(house_width / 2)],
            [house_width / 2, panel_y_start, roof_z(house_width / 2)]
        ])

        # Część lewej połaci po lewej stronie otworu
        left_roof_polygons.append([
            [panel_x_start, panel_y_end, roof_z(panel_x_start)],
            [house_width / 2, panel_y_end, roof_z(house_width / 2)],
            [house_width / 2, house_depth, roof_z(house_width / 2)],
            [0, house_depth, roof_z(0)],
            [0, panel_y_end, roof_z(0)]
        ])

        # Część lewej połaci poniżej otworu
        left_roof_polygons.append([
            [panel_x_start, panel_y_start, roof_z(panel_x_start)],
            [panel_x_start, panel_y_end, roof_z(panel_x_start)],
            [0, panel_y_end, roof_z(0)],
            [0, panel_y_start, roof_z(0)]
        ])

        # Część lewej połaci po prawej stronie otworu
        left_roof_polygons.append([
            [0, 0, roof_z(0)],
            [house_width / 2, 0, roof_z(house_width / 2)],
            [house_width / 2, panel_y_start, roof_z(house_width / 2)],
            [0, panel_y_start, roof_z(0)]
        ])

        # Rysowanie lewej połaci dachu z dziurą
        for polygon in left_roof_polygons:
            self.ax.add_collection3d(Poly3DCollection([polygon],
                                                      facecolors='#A52A2A',
                                                      alpha=1.0))

        # Definiowanie prawej połaci dachu jako jednolitej
        right_roof_polygon = [
            [house_width / 2, 0, roof_z(house_width / 2)],
            [house_width, 0, roof_z(house_width)],
            [house_width, house_depth, roof_z(house_width)],
            [house_width / 2, house_depth, roof_z(house_width / 2)]
        ]

        # Rysowanie prawej połaci dachu jako jednolitej powierzchni
        self.ax.add_collection3d(Poly3DCollection([right_roof_polygon],
                                                  facecolors='#A52A2A', alpha=1.0))

        # Rysowanie przedniego trójkąta (fasady)
        front_triangle = [
            [0, 0, house_height],
            [house_width / 2, 0, house_height + roof_height],
            [house_width, 0, house_height]
        ]
        self.ax.add_collection3d(Poly3DCollection([front_triangle],
                                                  facecolors='#A52A2A', alpha=1.0))

        # Rysowanie tylnego trójkąta (fasady)
        back_triangle = [
            [0, house_depth, house_height],
            [house_width / 2, house_depth, house_height + roof_height],
            [house_width, house_depth, house_height]
        ]
        self.ax.add_collection3d(Poly3DCollection([back_triangle],
                                                  facecolors='#A52A2A', alpha=1.0))

        # Opcjonalne rysowanie krawędzi otworu
        hole_edges = [
            [panel_x_start, panel_y_start, roof_z(panel_x_start)],
            [panel_x_end, panel_y_start, roof_z(panel_x_end)],
            [panel_x_end, panel_y_end, roof_z(panel_x_end)],
            [panel_x_start, panel_y_end, roof_z(panel_x_start)],
            [panel_x_start, panel_y_start, roof_z(panel_x_start)]
        ]
        self.ax.plot([v[0] for v in hole_edges],
                     [v[1] for v in hole_edges],
                     [v[2] for v in hole_edges],
                     color='black')

        # Ukrywanie osi
        self.ax.set_axis_off()

        # Ustawienie zakresu osi, aby obejmowały całą konstrukcję
        self.ax.set_xlim([-1 - window_thickness, house_width + 1 + window_thickness])
        self.ax.set_ylim([-1 - window_thickness, house_depth + 1 + window_thickness])
        self.ax.set_zlim([-1, house_height + roof_height + 1])

        # Ustawienie odpowiedniego kąta widzenia
        self.ax.view_init(elev=self.current_elev, azim=self.current_azim)

        # Dodanie obsługi przybliżania i oddalania za pomocą scrolla myszy
        self.fig.canvas.mpl_connect('scroll_event', self.obsluga_przewijania)

        # Odświeżanie płótna
        self.canvas.draw()

    # Metoda do automatycznego obracania wizualizacji
    def auto_rotate(self):
        if self.ax is None or not isinstance(self.ax, Axes3D):
            return  # Uniknięcie błędów, jeśli nie mamy aktywnej osi 3D

        if not self.is_rotating_manually:
            self.current_azim += self.auto_rotate_speed
            if self.current_azim >= 360:
                self.current_azim = 0

            self.ax.view_init(elev=self.current_elev, azim=self.current_azim)
            self.canvas.draw_idle()

        # Ponowne wywołanie auto_rotate po 100 ms
        self.auto_rotate_id = self.root.after(100, self.auto_rotate)

    # Metoda do włączenia ręcznego obrotu
    def on_mouse_press(self, event):
        self.is_rotating_manually = True

    # Metoda do wznowienia automatycznego obrotu po zwolnieniu myszy
    def on_mouse_release(self, event):
        self.current_azim = self.ax.azim
        self.current_elev = self.ax.elev

        # Odczekaj określony czas i wznow automatyczny obrót, jeśli użytkownik przestanie obracać
        self.fig.canvas.get_tk_widget().after(self.auto_rotate_delay, self.reset_manual_rotation)

    # Metoda monitorująca ruch myszki
    def on_mouse_move(self, event):
        if self.ax is None or not isinstance(self.ax, Axes3D):
            return  # Nie wykonuj, jeśli oś nie istnieje lub nie jest 3D
        if event.button is not None:
            self.is_rotating_manually = True
        self.ogranicz_obrot(event)

    # Wznowienie automatycznego obrotu po czasie bezczynności
    def reset_manual_rotation(self):
        self.is_rotating_manually = False

    # Ograniczanie obrotu, aby nie można było zajrzeć pod dom
    def ogranicz_obrot(self, event=None):
        if not isinstance(self.ax, Axes3D):  # Sprawdzenie, czy self.ax to oś 3D
            return

        current_elev = self.ax.elev
        current_azim = self.ax.azim

        # Ograniczenie elewacji do wartości między 15 a 90 stopni (blokowanie widoku od dołu)
        if current_elev < 15:
            current_elev = 15
        elif current_elev > 90:
            current_elev = 90

        self.ax.view_init(elev=current_elev, azim=current_azim)
        self.canvas.draw_idle()

    def obsluga_przewijania(self, event):
        base_scale = 1.1
        if event.button == 'up':
            scale_factor = 1 / base_scale
        elif event.button == 'down':
            scale_factor = base_scale
        else:
            return

        # Pobieranie aktualnych limitów osi
        xlim = self.ax.get_xlim3d()
        ylim = self.ax.get_ylim3d()
        zlim = self.ax.get_zlim3d()

        # Obliczanie środków osi
        xcenter = np.mean(xlim)
        ycenter = np.mean(ylim)
        zcenter = np.mean(zlim)

        # Aktualizacja limitów osi z uwzględnieniem marginesu
        new_xlim = [xcenter + (x - xcenter) * scale_factor for x in xlim]
        new_ylim = [ycenter + (y - ycenter) * scale_factor for y in ylim]
        new_zlim = [zcenter + (z - zcenter) * scale_factor for z in zlim]

        self.ax.set_xlim3d(new_xlim)
        self.ax.set_ylim3d(new_ylim)
        self.ax.set_zlim3d(new_zlim)

        # Zastosowanie ograniczenia obrotu
        self.ogranicz_obrot(event)

        # Odświeżenie rysunku
        self.canvas.draw_idle()

    def update_view(self):
        self.clear_current_view()
        if self.current_view == 'Wykres słupkowy':
            self.pokaz_wykres_slupkowy()
        elif self.current_view == 'Wizualizacja 3D':
            self.pokaz_wizualizacje_3d()

    def ustaw_naswietlenie(self):
        # Okno do ustawiania naświetlenia miesięcznego
        self.naswietlenie_window = tk.Toplevel(self.root)
        self.naswietlenie_window.title(self.texts['monthly sunshine hours'])

        # Ustawienie konkretnej wielkości okna i zablokowanie możliwości zmiany rozmiaru
        self.naswietlenie_window.geometry("280x380")
        self.naswietlenie_window.resizable(False, False)

        # Wyśrodkowanie okna na ekranie
        window_width = self.naswietlenie_window.winfo_reqwidth()
        window_height = self.naswietlenie_window.winfo_reqheight()
        position_right = int(self.naswietlenie_window.winfo_screenwidth() / 2 - window_width / 2)
        position_down = int(self.naswietlenie_window.winfo_screenheight() / 2 - window_height / 2)
        self.naswietlenie_window.geometry("+{}+{}".format(position_right, position_down))

        # Utworzenie ramki do wyśrodkowania elementów
        frame = tk.Frame(self.naswietlenie_window)
        frame.place(relx=0.5, rely=0.5, anchor='center')

        miesiace = self.texts['months']
        self.naswietlenie_entries = {}
        for i, miesiac in enumerate(miesiace):
            ttk.Label(frame, text=miesiac).grid(row=i, column=0, sticky='w')
            entry = ttk.Entry(frame)
            entry.grid(row=i, column=1)
            self.naswietlenie_entries[i + 1] = entry  # zmieniamy klucz na numer miesiąca

            # Wypełnienie pól wprowadzania danymi, jeśli są dostępne
            if i + 1 in self.naswietlenie_miesieczne:
                entry.insert(0, self.naswietlenie_miesieczne[i + 1])

        button_frame = tk.Frame(frame)
        button_frame.grid(row=len(miesiace), column=0, columnspan=2, pady=10)
        ttk.Button(button_frame, text=self.texts['save'], command=self.zapisz_naswietlenie).pack(side='left', padx=10)
        ttk.Button(button_frame, text=self.texts['delete'], command=self.kasuj_naswietlenie).pack(side='right', padx=10)

    def kasuj_naswietlenie(self):
        if messagebox.askyesno(self.texts['confirmation'], self.texts['delete_all_data'],
                               parent=self.naswietlenie_window):
            self.naswietlenie_miesieczne.clear()
            for entry in self.naswietlenie_entries.values():
                entry.delete(0, tk.END)

    def ustaw_temperatury(self):
        # Okno do ustawiania temperatury otoczenia dla każdego miesiąca
        self.temperatura_window = tk.Toplevel(self.root)
        self.temperatura_window.title(self.texts['monthly_temperatures'])

        # Ustawienie konkretnej wielkości okna i zablokowanie możliwości zmiany rozmiaru
        self.temperatura_window.geometry("280x380")
        self.temperatura_window.resizable(False, False)

        # Wyśrodkowanie okna na ekranie
        window_width = self.temperatura_window.winfo_reqwidth()
        window_height = self.temperatura_window.winfo_reqheight()
        position_right = int(self.temperatura_window.winfo_screenwidth() / 2 - window_width / 2)
        position_down = int(self.temperatura_window.winfo_screenheight() / 2 - window_height / 2)
        self.temperatura_window.geometry("+{}+{}".format(position_right, position_down))

        # Utworzenie ramki do wyśrodkowania elementów
        frame = tk.Frame(self.temperatura_window)
        frame.place(relx=0.5, rely=0.5, anchor='center')

        miesiace = self.texts['months']
        self.temperatura_entries = {}
        for i, miesiac in enumerate(miesiace):
            ttk.Label(frame, text=miesiac).grid(row=i, column=0, sticky='w')
            entry = ttk.Entry(frame)
            entry.grid(row=i, column=1)
            self.temperatura_entries[i + 1] = entry

            # Wypełnienie pól wprowadzania danymi, jeśli są dostępne
            if i + 1 in self.temperatura_miesieczna:
                entry.insert(0, self.temperatura_miesieczna[i + 1])

        button_frame = tk.Frame(frame)
        button_frame.grid(row=len(miesiace), column=0, columnspan=2, pady=10)
        ttk.Button(button_frame, text=self.texts['save'], command=self.zapisz_temperatury).pack(side='left', padx=10)
        ttk.Button(button_frame, text=self.texts['delete'], command=self.kasuj_temperatury).pack(side='right', padx=10)

    def kasuj_temperatury(self):
        if messagebox.askyesno(self.texts['confirmation'], self.texts['delete_all_data'],
                               parent=self.temperatura_window):
            self.temperatura_miesieczna.clear()
            for entry in self.temperatura_entries.values():
                entry.delete(0, tk.END)

    def zapisz_temperatury(self):
        try:
            for miesiac, entry in self.temperatura_entries.items():
                self.temperatura_miesieczna[miesiac] = float(entry.get())
            messagebox.showinfo(self.texts['success'], self.texts['monthly_temperatures_saved'])
            self.temperatura_window.destroy()
        except ValueError:
            messagebox.showerror(self.texts['error'], self.texts['enter_valid_numeric_values'])

    def zapisz_naswietlenie(self):
        try:
            for miesiac, entry in self.naswietlenie_entries.items():
                self.naswietlenie_miesieczne[miesiac] = float(entry.get())
            messagebox.showinfo(self.texts['success'], self.texts['monthly_illumination_saved'])
            self.naswietlenie_window.destroy()
        except ValueError:
            messagebox.showerror(self.texts['error'], self.texts['enter_valid_numeric_values'])

    def oblicz_wytworzona_energie(self, moc_instalacji, temperatura_otoczenia,
                                  naswietlenie, kąt_nachylenia, azymut, zacienienie, zanieczyszczenia):
        # Standardowy współczynnik wydajności systemu (PR)
        PR_standard = 0.6

        # Współczynnik temperaturowy
        Temperature_factor = 1 - 0.005 * (temperatura_otoczenia - 25)


        # Optymalne ustawienie to kąt nachylenia 30°
        tilt_loss = abs(kąt_nachylenia - 30) * 0.005
        SF = 1 - tilt_loss

        # Ograniczenie SF do zakresu 0 - 1
        SF = max(min(SF, 1), 0)

        # Całkowity PR
        PR = PR_standard * Temperature_factor * SF * (1 - zacienienie / 100) * (1 - zanieczyszczenia / 100)

        # Obliczanie energii (kWh)
        energia = (moc_instalacji * naswietlenie * PR) / 1000
        return energia

    def oblicz(self):
        try:
            liczba_paneli = int(self.entries["Liczba paneli"].get())
            if liczba_paneli <= 0:
                raise ValueError('invalid_number_of_panels')

            moc_jednego_panelu = float(self.entries["Moc jednego panelu (W)"].get())
            if moc_jednego_panelu <= 0:
                raise ValueError('invalid_panel_power')

            moc_instalacji = moc_jednego_panelu * liczba_paneli  # w Watach

            kąt_nachylenia = float(self.entries["Kąt nachylenia (°)"].get())
            if not 0 <= kąt_nachylenia <= 90:
                raise ValueError('invalid_tilt_angle')

            azymut = float(self.entries["Azymut (°)"].get())
            if not -180 <= azymut <= 180:
                raise ValueError('invalid_azimuth_angle')

            zacienienie = float(self.entries["Zacienienie (%)"].get())
            if not 0 <= zacienienie <= 100:
                raise ValueError('invalid_shading_percentage')

            zanieczyszczenia = float(self.entries["Zanieczyszczenia (%)"].get())
            if not 0 <= zanieczyszczenia <= 100:
                raise ValueError('invalid_soiling_percentage')

            # Sprawdzenie temperatury otoczenia
            if self.temperatura_miesieczna:
                use_monthly_temperatures = True
            else:
                temperatura_otoczenia = float(self.entries["Temperatura otoczenia (°C)"].get())
                if not -50 <= temperatura_otoczenia <= 60:
                    raise ValueError('invalid_ambient_temperature')
                use_monthly_temperatures = False

            # Sprawdzenie nasłonecznienia
            if self.naswietlenie_miesieczne:
                use_monthly_naswietlenie = True
            else:
                naswietlenie_roczne = float(self.entries["Liczba godzin słonecznych"].get())
                if not 0 <= naswietlenie_roczne <= 4380:
                    raise ValueError('invalid_sunlight_hours')
                use_monthly_naswietlenie = False

            energia_roczna = 0
            energia_miesieczna = []
            for i in range(1, 13):
                if use_monthly_naswietlenie:
                    naswietlenie = self.naswietlenie_miesieczne.get(i)
                    if naswietlenie is None or naswietlenie < 0:
                        raise ValueError('invalid_monthly_insolation')
                else:
                    naswietlenie = naswietlenie_roczne / 12

                if use_monthly_temperatures:
                    temperatura = self.temperatura_miesieczna.get(i)
                    if temperatura is None or not -50 <= temperatura <= 60:
                        raise ValueError('invalid_monthly_temperature')
                else:
                    temperatura = temperatura_otoczenia

                energia = self.oblicz_wytworzona_energie(
                    moc_instalacji, temperatura, naswietlenie, kąt_nachylenia, azymut,
                    zacienienie, zanieczyszczenia
                )

                if energia < 0:
                    raise ValueError('negative_monthly_energy')

                energia_roczna += energia
                energia_miesieczna.append(energia)

            if energia_roczna <= 0:
                raise ValueError('negative_annual_energy')

            # Wyświetl całkowitą energię roczną
            self.result_label.config(text=f"{self.texts['generated_energy']}: {energia_roczna:.2f} kWh")

            self.miesiace = self.texts['months']
            self.energia_miesieczna = energia_miesieczna

            self.pokaz_wykres_slupkowy()

        except ValueError as e:
            # Wyświetl komunikat o błędzie
            self.result_label.config(text=self.texts.get(str(e), "Wprowadzono nieprawidłowe dane"))

    def zapisz_wykres(self):
        # Podpowiedź dla nazwy pliku
        lista_plikow = glob.glob('wykres*.png')
        numery_plikow = [int(plik.split('.')[0].split('wykres')[1]) for plik in lista_plikow if
                         plik.split('.')[0].split('wykres')[1].isdigit()]
        nowy_numer = max(numery_plikow) + 1 if numery_plikow else 1
        propozycja_nazwy = f'wykres{nowy_numer}.png'

        # Pytaj użytkownika o nazwę pliku
        nazwa_pliku = askstring(self.texts['file_name_prompt'], self.texts['enter_file_name'],
                                initialvalue=propozycja_nazwy, texts=self.texts)

        # Zapisz wykres do pliku, jeśli nazwa_pliku nie jest None
        if nazwa_pliku is not None:
            self.fig.savefig(nazwa_pliku)
            messagebox.showinfo(self.texts['info'], f"{self.texts['chart_saved']} {nazwa_pliku} {self.texts['in_this_folder']}")

    def zapisz_ustawienia(self):
        # Podpowiedź dla nazwy pliku
        lista_plikow = glob.glob('ustawienia*.xlsx')
        numery_plikow = [int(plik.split('.')[0].split('ustawienia')[1]) for plik in lista_plikow if
                         plik.split('.')[0].split('ustawienia')[1].isdigit()]
        nowy_numer = max(numery_plikow) + 1 if numery_plikow else 1
        propozycja_nazwy = f'ustawienia{nowy_numer}.xlsx'

        # Pytaj użytkownika o nazwę pliku
        nazwa_pliku = askstring(self.texts['file_name_prompt'], self.texts['enter_file_name'],
                                initialvalue=propozycja_nazwy, texts=self.texts)

        # Zapisz ustawienia do pliku, jeśli nazwa_pliku nie jest None
        if nazwa_pliku is not None:
            df = pd.DataFrame({'naswietlenie_miesieczne': list(self.naswietlenie_miesieczne.values()),
                               'temperatura_miesieczna': list(self.temperatura_miesieczna.values())},
                              index=list(self.naswietlenie_miesieczne.keys()))
            df.to_excel(nazwa_pliku)
            messagebox.showinfo(self.texts['info'],
                                f"{self.texts['settings_saved']} {nazwa_pliku} {self.texts['in_this_folder']}")

    def zapisz_dane(self):
        # Pobieranie listy plików pasujących do wzorca 'dane*.xlsx'
        lista_plikow = glob.glob('dane*.xlsx')

        # Przetwarzanie listy plików, aby znaleźć numery plików w formacie 'daneX.xlsx'
        numery_plikow = [
            int(plik.split('.')[0].split('dane')[1])
            for plik in lista_plikow
            if 'dane' in plik.split('.')[0] and plik.split('.')[0].split('dane')[1].isdigit()
        ]

        # Ustalenie nowego numeru pliku
        nowy_numer = max(numery_plikow) + 1 if numery_plikow else 1
        propozycja_nazwy = f'dane{nowy_numer}.xlsx'

        # Pytanie użytkownika o nazwę pliku
        nazwa_pliku = askstring(self.texts['file_name_prompt'], self.texts['enter_file_name'],
                                initialvalue=propozycja_nazwy, texts=self.texts)

        # Zapis danych do pliku, jeśli użytkownik poda nazwę
        if nazwa_pliku is not None:
            dane = {label: entry.get() for label, entry in self.entries.items()}
            df = pd.DataFrame(dane, index=[0])
            df.to_excel(nazwa_pliku, index=False)
            messagebox.showinfo(self.texts['info'],
                                f"{self.texts['data_saved']} {nazwa_pliku} {self.texts['in_this_folder']}")

    def importuj_wykres(self):
        # Pozwól użytkownikowi wybrać plik do importu
        nazwa_pliku = filedialog.askopenfilename(filetypes=[("Pliki PNG", "*.png")])
        if nazwa_pliku:
            self.ax.clear()
            img = plt.imread(nazwa_pliku)
            self.ax.imshow(img, aspect='auto')
            self.ax.axis('off')
            self.canvas.draw()
            messagebox.showinfo(self.texts["success"], self.texts['chart_imported'])

    def sprawdz_plik(self, df, oczekiwane_kolumny):
        if len(df.columns) != len(oczekiwane_kolumny):
            raise ValueError(self.texts['wrong_number_of_columns'])

        for kolumna in oczekiwane_kolumny:
            if kolumna not in df.columns:
                raise ValueError(f"{self.texts['missing_expected_column']}: {kolumna}")

        return True

    def wybierz_i_importuj_ustawienia(self):
        try:
            nazwa_pliku = filedialog.askopenfilename(filetypes=[(self.texts["excel_files"], "*.xlsx")])
            if nazwa_pliku:
                df = pd.read_excel(nazwa_pliku, index_col=0)

                # Sprawdź, czy plik jest poprawny
                if not self.sprawdz_plik(df, ['naswietlenie_miesieczne', 'temperatura_miesieczna']):
                    raise ValueError(self.texts["file_is_incorrect_or_damaged"])

                self.naswietlenie_miesieczne = {i + 1: v for i, v in enumerate(df['naswietlenie_miesieczne'].tolist())}
                self.temperatura_miesieczna = {i + 1: v for i, v in enumerate(df['temperatura_miesieczna'].tolist())}

                # Aktualizacja interfejsu użytkownika z nowymi wartościami
                for i, miesiac in enumerate(self.texts['months'], start=1):
                    if hasattr(self, 'naswietlenie_entries') and i in self.naswietlenie_entries and \
                            self.naswietlenie_entries[i].winfo_exists() and i in self.naswietlenie_miesieczne:
                        self.naswietlenie_entries[i].delete(0, tk.END)
                        self.naswietlenie_entries[i].insert(0, self.naswietlenie_miesieczne[i])

                    if hasattr(self, 'temperatura_entries') and i in self.temperatura_entries and \
                            self.temperatura_entries[i].winfo_exists() and i in self.temperatura_miesieczna:
                        self.temperatura_entries[i].delete(0, tk.END)
                        self.temperatura_entries[i].insert(0, self.temperatura_miesieczna[i])
                messagebox.showinfo(self.texts["success"], self.texts["settings_imported_successfully"])
        except Exception as e:
            messagebox.showerror(self.texts["error"], str(e))

    def wybierz_i_importuj_dane(self):
        try:
            # Pozwól użytkownikowi wybrać plik do importu
            nazwa_pliku = filedialog.askopenfilename(filetypes=[(self.texts["excel_files"], "*.xlsx")])
            if nazwa_pliku:
                df = pd.read_excel(nazwa_pliku)

                # Sprawdź, czy plik jest poprawny
                if not self.sprawdz_plik(df, list(self.entries.keys())):
                    raise ValueError(self.texts["file_is_incorrect_or_damaged"])

                dane = df.to_dict(orient='records')[0]

                # Aktualizacja interfejsu użytkownika z nowymi wartościami
                for label, wartosc in dane.items():
                    if label in self.entries:
                        self.entries[label].delete(0, tk.END)
                        self.entries[label].insert(0, wartosc)
                messagebox.showinfo(self.texts["success"], self.texts["data_imported_successfully"])
        except Exception as e:
            messagebox.showerror(self.texts["error"], str(e))

    def pokaz_pomoc(self, event=None):
        """Wyświetlanie pomocy w głównym obszarze aplikacji."""
        # Sprawdź, czy plik pomocy jest już wyświetlany
        if hasattr(self, 'help_displayed') and self.help_displayed:
            return  # Jeśli tak, zakończ funkcję

        self.clear_current_view()
        self.chart_frame.pack_forget()

        # Utwórz główny kontener pomocy
        self.help_frame = tk.Frame(self.chart_container, bg='white')
        self.help_frame.pack(fill='both', expand=True)

        # Utwórz przewijaną ramkę
        canvas = tk.Canvas(self.help_frame, bg='white', highlightthickness=0)
        scrollbar = tk.Scrollbar(self.help_frame, orient="vertical", command=canvas.yview)
        text_frame = tk.Frame(canvas, bg='white')

        # Konfiguracja przewijania
        text_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        # Główny widget tekstowy
        main_text = tk.Text(text_frame, wrap=tk.WORD, bg='white', padx=20, pady=20,
                            font=('Arial', 16), spacing3=10, borderwidth=0)
        main_text.pack(fill='both', expand=True)
        main_text.images = []

        # Ładowanie zawartości
        help_file_name = "Pomoc/Pomoc.docx" if self.language == 'polski' else "Pomoc/Help.docx"
        help_file = self.get_resource_path(help_file_name)
        try:
            self.doc = Document(help_file)
            self._process_docx_content(self.doc, main_text)
        except Exception as e:
            main_text.insert(tk.END, f"\n\nBłąd ładowania pomocy: {str(e)}", 'error')

        # Dynamiczne dostosowanie szerokości
        def update_width(event):
            canvas.itemconfig(text_frame_id, width=event.width)

        canvas.bind("<Configure>", update_width)

        # Układ elementów
        text_frame_id = canvas.create_window((0, 0), window=text_frame, anchor='nw')
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')

        self.current_view = 'help'
        self.help_displayed = True  # Ustaw flagę, że plik pomocy jest wyświetlany

    def _process_docx_content(self, doc, text_widget):
        """Przetwarzanie zawartości dokumentu DOCX"""
        for paragraph in doc.paragraphs:
            self._add_paragraph(paragraph, text_widget)
            for run in paragraph.runs:
                self._check_for_images(run, text_widget)

        for table in doc.tables:
            self._add_table(table, text_widget)

    def _check_for_images(self, run, text_widget):
        """Sprawdź czy run zawiera obraz"""
        try:
            if run._element.xpath('.//wp:inline'):
                self._add_image(run, text_widget)
        except Exception as e:
            print(f"Błąd przetwarzania obrazu: {str(e)}")

    def _add_image(self, run, text_widget):
        """Dodawanie obrazu z run"""
        try:
            drawing = run._element.xpath('.//wp:inline')[0]
            blip = drawing.xpath('.//a:blip/@r:embed')[0]
            image_part = run.part.related_parts[blip]
            image_bytes = image_part.blob

            image = Image.open(io.BytesIO(image_bytes))
            image.thumbnail((600, 600))

            photo_image = ImageTk.PhotoImage(image)
            text_widget.image_create(tk.END, image=photo_image)
            text_widget.images.append(photo_image)
            text_widget.insert(tk.END, "\n")
        except Exception as e:
            text_widget.insert(tk.END, f"\n[BRAK OBRAZU: {str(e)}]\n")

    def _add_paragraph(self, paragraph, text_widget):
        """Dodawanie paragrafu tekstu"""
        text_widget.insert(tk.END, paragraph.text + "\n\n")

    def _add_table(self, table, text_widget):
        """Dodawanie tabeli"""
        text_widget.insert(tk.END, "\nTABELA:\n")
        for row in table.rows:
            for cell in row.cells:
                text_widget.insert(tk.END, f"{cell.text} | ")
            text_widget.insert(tk.END, "\n")
        text_widget.insert(tk.END, "\n")

    def wyjdz(self):
        if messagebox.askokcancel(self.texts['exit'], self.texts['exit_confirmation']):
            self.root.destroy()

    def clear_current_view(self):
        # Anulowanie automatycznego obracania, jeśli jest aktywne
        if self.auto_rotate_id:
            self.root.after_cancel(self.auto_rotate_id)
            self.auto_rotate_id = None

        # Usunięcie eventów myszki, aby nie były powiązane ze starymi obiektami
        if self.fig and self.fig.canvas:
            self.fig.canvas.mpl_disconnect(self.fig.canvas.mpl_connect('motion_notify_event', self.on_mouse_move))
            self.fig.canvas.mpl_disconnect(self.fig.canvas.mpl_connect('button_press_event', self.on_mouse_press))
            self.fig.canvas.mpl_disconnect(self.fig.canvas.mpl_connect('button_release_event', self.on_mouse_release))

        # Usunięcie osi, jeśli istnieją
        if self.ax is not None:
            self.ax.remove()
            self.ax = None

        # Wyczyszczenie figury, ale bez jej całkowitego usuwania
        if self.fig:
            self.fig.clf()

        # Odświeżenie canvas, jeśli istnieje
        if self.canvas:
            self.canvas.draw_idle()


if __name__ == "__main__":
    root = tk.Tk()
    app = MojaAplikacja(root)
    root.mainloop()